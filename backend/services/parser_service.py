import os
import uuid
import logging
from PyPDF2 import PdfReader
from docx import Document
from fastapi import UploadFile, HTTPException

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB


async def save_upload_file(file: UploadFile, upload_dir: str = None) -> str:
    """Save uploaded file and return the file path."""
    if upload_dir is None:
        upload_dir = os.getenv("UPLOAD_DIR", "../uploads")
    os.makedirs(upload_dir, exist_ok=True)

    content = await file.read()

    # Enforce file size limit
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail="File too large. Maximum size is 10 MB."
        )

    # Sanitize filename: use UUID to prevent path traversal
    original_ext = os.path.splitext(file.filename or "unknown.pdf")[1].lower()
    if original_ext not in [".pdf", ".docx", ".doc"]:
        original_ext = ".pdf"
    safe_filename = f"{uuid.uuid4().hex}{original_ext}"
    file_path = os.path.join(upload_dir, safe_filename)

    # Verify the resolved path stays within the upload directory
    abs_upload_dir = os.path.abspath(upload_dir)
    abs_file_path = os.path.abspath(file_path)
    if not abs_file_path.startswith(abs_upload_dir):
        raise HTTPException(status_code=400, detail="Invalid file path")

    with open(file_path, "wb") as f:
        f.write(content)

    return file_path


def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {type(e).__name__}")


def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {type(e).__name__}")


def parse_resume(file_path: str) -> str:
    """Parse resume based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return parse_docx(file_path)
    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Please upload PDF or DOCX."
        )


def cleanup_file(file_path: str) -> None:
    """Remove temporary file after processing."""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except Exception as e:
        logger.warning(f"Failed to cleanup uploaded file: {type(e).__name__}")
